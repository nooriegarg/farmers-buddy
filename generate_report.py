from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# Create document
doc = Document()

# Set page size to A4
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)

# Helper functions
def set_cell_bg(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def add_heading(doc, text, level=1, color_hex='1B5E20'):
    heading = doc.add_heading(text, level=level)
    run = heading.runs[0] if heading.runs else heading.add_run(text)
    run.font.color.rgb = RGBColor.from_string(color_hex)
    return heading

def add_footer_page_numbers(doc):
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

# Add footer
add_footer_page_numbers(doc)

# ============ TITLE PAGE ============
title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
for _ in range(4):
    doc.add_paragraph()

title_run = title_para.add_run('FARMERS BUDDY')
title_run.font.size = Pt(36)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(27, 94, 32)

subtitle = doc.add_paragraph('Agriculture Support Platform')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(18)
subtitle.runs[0].font.color.rgb = RGBColor(46, 125, 50)

doc.add_paragraph()

report_type = doc.add_paragraph('Final Year Project Report')
report_type.alignment = WD_ALIGN_PARAGRAPH.CENTER
report_type.runs[0].font.size = Pt(14)
report_type.runs[0].font.bold = True

doc.add_paragraph()
doc.add_paragraph()

# Student details
fields = [
    'Student Name: ___________________________',
    'Roll Number: ___________________________',
    'College Name: ___________________________',
    'Department: Computer Science / IT',
    'Academic Year: 2025-2026',
    'Guide / Mentor: ___________________________',
]
for field in fields:
    p = doc.add_paragraph(field)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(12)

doc.add_page_break()

# ============ SECTION 1: PROBLEM STATEMENT ============
add_heading(doc, '1. Problem Statement')
doc.add_paragraph(
    'Agriculture is the backbone of the Indian economy, yet farmers continue to face significant challenges '
    'in accessing timely and accurate guidance from agricultural officers. The traditional methods of '
    'communication — physical visits to government offices, phone calls, or relying on intermediaries — '
    'are slow, inefficient, and inaccessible to farmers in remote areas.'
)
doc.add_paragraph(
    'Farmers often have urgent queries regarding crop diseases, fertilizer usage, weather impacts, and '
    'government schemes. Delays in receiving expert advice can lead to crop damage and economic losses. '
    'There is a clear need for a dedicated digital platform that bridges the communication gap between '
    'farmers and agricultural officers.'
)
doc.add_paragraph(
    'Farmers Buddy is developed to address this problem by providing a centralized, role-based web '
    'application where farmers can digitally submit their agriculture-related queries, and officers can '
    'manage, review, and respond to these queries in a timely and organized manner.'
)

# ============ SECTION 2: OBJECTIVES ============
doc.add_page_break()
add_heading(doc, '2. Objectives')
objectives = [
    'Provide farmers with a user-friendly digital platform to submit agriculture-related queries anytime, anywhere.',
    'Enable agricultural officers to manage, filter, search, and respond to farmer queries efficiently.',
    'Implement a secure, role-based authentication system distinguishing Farmer and Officer roles.',
    'Build a clean, layered full-stack architecture using industry-standard technologies.',
    'Ensure persistent and reliable data storage using MySQL with JPA/Hibernate ORM.',
    'Create responsive UI using Tailwind CSS for accessibility on multiple devices.',
    'Follow RESTful API design principles for clean frontend-backend communication.',
]
for obj in objectives:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(obj)

# ============ SECTION 3: TECHNOLOGY STACK ============
doc.add_page_break()
add_heading(doc, '3. Technology Stack')
doc.add_paragraph(
    'The following technologies were carefully selected based on industry standards, scalability, '
    'and suitability for a full-stack web application:'
)

tech_table = doc.add_table(rows=1, cols=3)
tech_table.style = 'Table Grid'
tech_table.alignment = WD_TABLE_ALIGNMENT.CENTER

header_cells = tech_table.rows[0].cells
for i, h in enumerate(['Layer', 'Technology', 'Purpose']):
    header_cells[i].text = h
    header_cells[i].paragraphs[0].runs[0].font.bold = True
    set_cell_bg(header_cells[i], '2E7D32')
    header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

tech_data = [
    ('Frontend', 'React.js', 'UI Development & Component-Based Architecture'),
    ('Frontend', 'Tailwind CSS', 'Styling & Responsive Design'),
    ('Frontend', 'Axios', 'HTTP Client for REST API Calls'),
    ('Frontend', 'React Router', 'Client-Side Routing & Navigation'),
    ('Backend', 'Spring Boot', 'REST API Framework & Application Server'),
    ('Backend', 'Spring Security', 'Authentication & Authorization'),
    ('Backend', 'Spring Data JPA', 'ORM & Database Operations'),
    ('Database', 'MySQL', 'Relational Database Management System'),
    ('Tools', 'VS Code', 'Frontend Development IDE'),
    ('Tools', 'IntelliJ IDEA', 'Backend Development IDE'),
    ('Tools', 'GitHub', 'Version Control & Collaboration'),
]

for i, (layer, tech, purpose) in enumerate(tech_data):
    row = tech_table.add_row()
    row.cells[0].text = layer
    row.cells[1].text = tech
    row.cells[2].text = purpose
    bg = 'E8F5E9' if i % 2 == 0 else 'FFFFFF'
    for cell in row.cells:
        set_cell_bg(cell, bg)

# ============ SECTION 4: SYSTEM ARCHITECTURE DIAGRAM ============
doc.add_page_break()
add_heading(doc, '4. System Architecture Diagram')
doc.add_paragraph(
    'The system follows a standard 3-tier client-server architecture. The frontend React application '
    'communicates with the Spring Boot backend via RESTful HTTP APIs. The backend processes business '
    'logic and interacts with the MySQL database through JPA/Hibernate ORM.'
)

arch_table = doc.add_table(rows=5, cols=5)
arch_table.style = 'Table Grid'
arch_table.alignment = WD_TABLE_ALIGNMENT.CENTER

arch_data = [
    ('FRONTEND\nLAYER', '<-> HTTP/JSON', 'REST API\nLAYER', '<-> Business Logic', 'BACKEND\nLAYER'),
    ('React.js\nTailwind CSS\nAxios\nReact Router', '', 'Controllers\nAuthController\nQueryController', '', 'Service Classes\nAuthService\nQueryService'),
    ('', '', '', '', ''),
    ('', '', 'JPA/Hibernate\nORM Layer', '<-> SQL Queries', 'MySQL\nDatabase'),
    ('', '', 'Data Access\nRepositories', '', 'Persistent\nStorage'),
]

colors = ['1565C0', 'ECEFF1', '2E7D32', 'ECEFF1', '4527A0']
text_colors = [True, False, True, False, True]

for r, row_data in enumerate(arch_data):
    for c, cell_text in enumerate(row_data):
        cell = arch_table.rows[r].cells[c]
        cell.text = cell_text
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if cell_text:
            run = p.runs[0] if p.runs else p.add_run(cell_text)
            run.font.bold = True
            run.font.size = Pt(9)
            if r == 0 and text_colors[c]:
                set_cell_bg(cell, colors[c])
                run.font.color.rgb = RGBColor(255, 255, 255)
            elif r == 0:
                set_cell_bg(cell, 'ECEFF1')
            else:
                set_cell_bg(cell, 'E8F5E9' if c in [0, 2, 4] else 'FFFFFF')

doc.add_paragraph()
viva_box = doc.add_paragraph()
viva_box.add_run('Viva Tip: ').font.bold = True
viva_box.add_run(
    'The system uses a 3-tier architecture: Presentation (React), Logic (Spring Boot), and Data (MySQL). '
    'Communication between layers uses REST APIs with JSON data format. Hibernate acts as an ORM to map '
    'Java objects to database tables automatically.'
)

# ============ SECTION 5: LAYERED ARCHITECTURE DIAGRAM ============
doc.add_page_break()
add_heading(doc, '5. Layered Architecture Diagram')
doc.add_paragraph(
    'The application follows a clean separation of concerns through distinct architectural layers. '
    'Each layer has a specific responsibility and communicates only with adjacent layers.'
)

layer_table = doc.add_table(rows=5, cols=1)
layer_table.style = 'Table Grid'
layer_table.alignment = WD_TABLE_ALIGNMENT.CENTER

layers = [
    ('PRESENTATION LAYER', 'React.js Components | Tailwind CSS | React Router | Axios HTTP Client', '1565C0'),
    ('API / CONTROLLER LAYER', 'AuthController | QueryController | REST Endpoints | Request/Response Mapping', '2E7D32'),
    ('BUSINESS LOGIC LAYER', 'AuthService | QueryService | Validation | Password Encoding (BCrypt)', '6A1B9A'),
    ('DATA ACCESS LAYER', 'UserRepository | QueryRepository | Spring Data JPA | JPQL Queries', 'E65100'),
    ('DATABASE LAYER', 'MySQL | Tables: users, queries | Persistent Storage | ACID Compliance', '37474F'),
]

for i, (title, details, color) in enumerate(layers):
    cell = layer_table.rows[i].cells[0]
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = p.add_run(title + '\n')
    title_run.font.bold = True
    title_run.font.size = Pt(11)
    title_run.font.color.rgb = RGBColor(255, 255, 255)
    detail_run = p.add_run(details)
    detail_run.font.size = Pt(9)
    detail_run.font.color.rgb = RGBColor(255, 255, 255)
    set_cell_bg(cell, color)

doc.add_paragraph()
viva_box2 = doc.add_paragraph()
viva_box2.add_run('Viva Tip: ').font.bold = True
viva_box2.add_run(
    'The layered architecture ensures loose coupling. The Presentation Layer never directly accesses '
    'the database. Every request passes through Controller -> Service -> Repository -> Database. '
    'This makes the code maintainable, testable, and scalable.'
)

# ============ SECTION 6: LOGIN WORKFLOW ============
doc.add_page_break()
add_heading(doc, '6. Login Workflow Diagram')
doc.add_paragraph(
    'The login process involves multiple layers from user input on the frontend to database validation '
    'on the backend, finally redirecting to the appropriate role-based dashboard.'
)

login_steps = [
    ('1', 'User', 'Opens Login Page in Browser'),
    ('2', 'React Frontend', 'User fills Email & Password -> Clicks Login'),
    ('3', 'Axios', 'POST /api/auth/login with credentials (JSON)'),
    ('4', 'AuthController', 'Receives request -> Calls AuthService'),
    ('5', 'AuthService', 'Validates credentials -> BCrypt password check'),
    ('6', 'UserRepository', 'Finds user by email in MySQL database'),
    ('7', 'MySQL', 'Returns user record with role information'),
    ('8', 'AuthService', 'Generates auth token / returns user role'),
    ('9', 'React Frontend', 'Stores token in localStorage'),
    ('10', 'React Router', 'Redirects to Farmer or Officer Dashboard based on role'),
]

flow_table = doc.add_table(rows=len(login_steps), cols=3)
flow_table.style = 'Table Grid'
flow_table.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, (step, component, action) in enumerate(login_steps):
    row = flow_table.rows[i]
    row.cells[0].text = f'Step {step}'
    row.cells[1].text = component
    row.cells[2].text = action
    bg = 'E8F5E9' if i % 2 == 0 else 'FFFFFF'
    set_cell_bg(row.cells[0], '2E7D32')
    row.cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    row.cells[0].paragraphs[0].runs[0].font.bold = True
    set_cell_bg(row.cells[1], bg)
    set_cell_bg(row.cells[2], bg)

doc.add_paragraph()
doc.add_paragraph(
    'Flow Summary: User -> Login Form -> Axios API Call -> AuthController -> AuthService -> '
    'UserRepository -> MySQL -> Token -> localStorage -> Role-Based Dashboard'
)

# ============ SECTION 7: QUERY SUBMISSION WORKFLOW ============
doc.add_page_break()
add_heading(doc, '7. Query Submission Workflow')

query_steps = [
    ('1', 'Farmer', 'Logged-in Farmer opens Farmer Dashboard'),
    ('2', 'React Frontend', 'Farmer fills Query Title & Description -> Submits Form'),
    ('3', 'Axios', 'POST /api/queries with query data and auth token'),
    ('4', 'QueryController', 'Receives POST request -> Validates input'),
    ('5', 'QueryService', 'Processes query -> Sets status to PENDING'),
    ('6', 'QueryRepository', 'Saves query object using Spring Data JPA'),
    ('7', 'MySQL', 'Inserts new record into QUERIES table'),
    ('8', 'QueryController', 'Returns success response (201 Created)'),
    ('9', 'React Frontend', 'Shows success notification to farmer'),
    ('10', 'Farmer Dashboard', 'Updates query list with newly submitted query'),
]

q_table = doc.add_table(rows=len(query_steps), cols=3)
q_table.style = 'Table Grid'
q_table.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, (step, component, action) in enumerate(query_steps):
    row = q_table.rows[i]
    row.cells[0].text = f'Step {step}'
    row.cells[1].text = component
    row.cells[2].text = action
    bg = 'E3F2FD' if i % 2 == 0 else 'FFFFFF'
    set_cell_bg(row.cells[0], '1565C0')
    row.cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    row.cells[0].paragraphs[0].runs[0].font.bold = True
    set_cell_bg(row.cells[1], bg)
    set_cell_bg(row.cells[2], bg)

doc.add_paragraph()
doc.add_paragraph(
    'Flow Summary: Farmer Dashboard -> Query Form -> Axios POST -> QueryController -> '
    'QueryService -> QueryRepository -> MySQL Database -> Success Response'
)

# ============ SECTION 8: OFFICER REPLY WORKFLOW ============
doc.add_page_break()
add_heading(doc, '8. Officer Reply Workflow')

officer_steps = [
    ('1', 'Officer', 'Logged-in Officer opens Officer Dashboard'),
    ('2', 'React Frontend', 'Dashboard loads all farmer queries via GET /api/queries'),
    ('3', 'Officer', 'Officer can filter queries by status or search by farmer name'),
    ('4', 'Officer', 'Officer selects a PENDING query and writes a reply'),
    ('5', 'Axios', 'PUT /api/queries/{id}/reply with reply text and auth token'),
    ('6', 'QueryController', 'Receives PUT request -> Calls QueryService'),
    ('7', 'QueryService', 'Updates officerReply field -> Changes status to REPLIED'),
    ('8', 'QueryRepository', 'Saves updated query to MySQL via JPA'),
    ('9', 'MySQL', 'Updates existing record in QUERIES table'),
    ('10', 'Farmer Dashboard', 'Farmer can now see the officer reply on their query'),
]

o_table = doc.add_table(rows=len(officer_steps), cols=3)
o_table.style = 'Table Grid'
o_table.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, (step, component, action) in enumerate(officer_steps):
    row = o_table.rows[i]
    row.cells[0].text = f'Step {step}'
    row.cells[1].text = component
    row.cells[2].text = action
    bg = 'F3E5F5' if i % 2 == 0 else 'FFFFFF'
    set_cell_bg(row.cells[0], '6A1B9A')
    row.cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    row.cells[0].paragraphs[0].runs[0].font.bold = True
    set_cell_bg(row.cells[1], bg)
    set_cell_bg(row.cells[2], bg)

doc.add_paragraph()
doc.add_paragraph(
    'Flow Summary: Officer Dashboard -> View Queries -> Reply Form -> PUT API -> '
    'QueryController -> QueryService -> Update Query -> MySQL -> Farmer Dashboard Updated'
)

# ============ SECTION 9: ROLE-BASED ACCESS ============
doc.add_page_break()
add_heading(doc, '9. Role-Based Access Workflow')
doc.add_paragraph(
    "Farmers Buddy implements strict role-based access control. After login, the user role "
    "determines which dashboard and features are accessible. React Router's ProtectedRoute component "
    "guards routes on the frontend, while Spring Security secures API endpoints on the backend."
)

role_table = doc.add_table(rows=8, cols=2)
role_table.style = 'Table Grid'
role_table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Headers
set_cell_bg(role_table.rows[0].cells[0], '1565C0')
set_cell_bg(role_table.rows[0].cells[1], '2E7D32')
role_table.rows[0].cells[0].text = 'FARMER ROLE'
role_table.rows[0].cells[1].text = 'OFFICER ROLE'
for cell in role_table.rows[0].cells:
    cell.paragraphs[0].runs[0].font.bold = True
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

farmer_actions = ['Register as FARMER', 'Login -> Farmer Dashboard', 'Submit New Query', 'View Own Queries', 'View Officer Replies', 'Cannot Access Officer Dashboard', 'Cannot Reply to Queries']
officer_actions = ['Register as OFFICER', 'Login -> Officer Dashboard', 'View ALL Farmer Queries', 'Filter by Status / Search', 'Reply to Farmer Queries', 'Update Query Status', 'Cannot Access Farmer Dashboard']

for i in range(1, 8):
    row = role_table.rows[i]
    row.cells[0].text = farmer_actions[i-1]
    row.cells[1].text = officer_actions[i-1]
    set_cell_bg(row.cells[0], 'E3F2FD')
    set_cell_bg(row.cells[1], 'E8F5E9')

doc.add_paragraph()
viva_role = doc.add_paragraph()
viva_role.add_run('Viva Tip: ').font.bold = True
viva_role.add_run(
    "Role-based access is implemented at two levels: (1) Frontend - React Router's ProtectedRoute "
    "component checks the user role from localStorage before rendering a page. (2) Backend - Spring Security "
    "uses @PreAuthorize or security config to restrict API endpoints by role."
)

# ============ SECTION 10: DATABASE ARCHITECTURE ============
doc.add_page_break()
add_heading(doc, '10. Database Architecture')
doc.add_paragraph(
    'The database consists of two primary tables: USERS and QUERIES. These tables are related through '
    'the farmerName field. The USERS table stores authentication and role information, while the '
    'QUERIES table stores all farmer queries and officer replies.'
)

# ER Diagram representation
er_table = doc.add_table(rows=3, cols=3)
er_table.style = 'Table Grid'
er_table.alignment = WD_TABLE_ALIGNMENT.CENTER

# USERS box
users_cell = er_table.rows[0].cells[0]
set_cell_bg(users_cell, '1565C0')
users_cell.text = ''
p = users_cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('USERS TABLE\n')
r.font.bold = True
r.font.color.rgb = RGBColor(255, 255, 255)
r.font.size = Pt(11)
r2 = p.add_run('PK: id\nname\nemail (UNIQUE)\npassword\nrole (FARMER/OFFICER)')
r2.font.color.rgb = RGBColor(255, 255, 255)
r2.font.size = Pt(9)

# Relation
rel_cell = er_table.rows[0].cells[1]
rel_cell.text = 'One\nFarmer\nHas Many\nQueries\n-------->'
rel_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
set_cell_bg(rel_cell, 'ECEFF1')

# QUERIES box
queries_cell = er_table.rows[0].cells[2]
set_cell_bg(queries_cell, '2E7D32')
queries_cell.text = ''
p2 = queries_cell.paragraphs[0]
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p2.add_run('QUERIES TABLE\n')
r3.font.bold = True
r3.font.color.rgb = RGBColor(255, 255, 255)
r3.font.size = Pt(11)
r4 = p2.add_run('PK: id\nfarmerName (FK reference)\ntitle\ndescription\nofficerReply\nstatus')
r4.font.color.rgb = RGBColor(255, 255, 255)
r4.font.size = Pt(9)

doc.add_paragraph()
doc.add_paragraph(
    'Relationship: One FARMER (user) can submit MANY QUERIES. The relationship is maintained through '
    'the farmerName field in the QUERIES table, which stores the name of the farmer who submitted the query. '
    'Spring Data JPA manages the object-relational mapping between Java entities and database tables.'
)

# ============ SECTION 11: DATABASE TABLES STRUCTURE ============
doc.add_page_break()
add_heading(doc, '11. Database Tables Structure')

add_heading(doc, '11.1 USERS Table', level=2)
users_table = doc.add_table(rows=1, cols=3)
users_table.style = 'Table Grid'
for i, h in enumerate(['Column Name', 'Data Type', 'Constraints']):
    users_table.rows[0].cells[i].text = h
    users_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
    set_cell_bg(users_table.rows[0].cells[i], '1565C0')
    users_table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

users_cols = [
    ('id', 'BIGINT', 'PRIMARY KEY, AUTO_INCREMENT'),
    ('name', 'VARCHAR(100)', 'NOT NULL'),
    ('email', 'VARCHAR(150)', 'UNIQUE, NOT NULL'),
    ('password', 'VARCHAR(255)', 'NOT NULL - BCrypt Hashed'),
    ('role', 'VARCHAR(20)', 'NOT NULL - Values: FARMER or OFFICER'),
]
for i, row_data in enumerate(users_cols):
    row = users_table.add_row()
    for j, val in enumerate(row_data):
        row.cells[j].text = val
        set_cell_bg(row.cells[j], 'E3F2FD' if i % 2 == 0 else 'FFFFFF')

doc.add_paragraph()
add_heading(doc, '11.2 QUERIES Table', level=2)
queries_table = doc.add_table(rows=1, cols=3)
queries_table.style = 'Table Grid'
for i, h in enumerate(['Column Name', 'Data Type', 'Constraints']):
    queries_table.rows[0].cells[i].text = h
    queries_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
    set_cell_bg(queries_table.rows[0].cells[i], '2E7D32')
    queries_table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

queries_cols = [
    ('id', 'BIGINT', 'PRIMARY KEY, AUTO_INCREMENT'),
    ('farmerName', 'VARCHAR(100)', 'NOT NULL - References farmer username'),
    ('title', 'VARCHAR(200)', 'NOT NULL'),
    ('description', 'TEXT', 'NOT NULL - Full query description'),
    ('officerReply', 'TEXT', 'NULLABLE - Populated when officer replies'),
    ('status', 'VARCHAR(20)', 'DEFAULT: PENDING | Values: PENDING / REPLIED'),
]
for i, row_data in enumerate(queries_cols):
    row = queries_table.add_row()
    for j, val in enumerate(row_data):
        row.cells[j].text = val
        set_cell_bg(row.cells[j], 'E8F5E9' if i % 2 == 0 else 'FFFFFF')

# ============ SECTION 12: FEATURES IMPLEMENTED ============
doc.add_page_break()
add_heading(doc, '12. Features Implemented')

feat_table = doc.add_table(rows=1, cols=2)
feat_table.style = 'Table Grid'
for i, h in enumerate(['Feature', 'Description']):
    feat_table.rows[0].cells[i].text = h
    feat_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
    set_cell_bg(feat_table.rows[0].cells[i], '37474F')
    feat_table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

features = [
    ('User Registration', 'Farmers and Officers can register with name, email, password, and role selection'),
    ('Secure Login / Authentication', 'JWT/token-based authentication using Spring Security and BCrypt password hashing'),
    ('Farmer Dashboard', 'Farmers can submit new queries, view all their submitted queries, and see officer replies'),
    ('Officer Dashboard', 'Officers can view all farmer queries, filter by status, search by farmer name, and reply'),
    ('Query Submission', 'Farmers submit agriculture queries with title and detailed description'),
    ('Query Management', 'Full CRUD operations - Create, Read, Update (reply) for queries'),
    ('Officer Reply System', 'Officers reply to queries which updates the officerReply field and status to REPLIED'),
    ('Role-Based Routing', 'React Router ProtectedRoute ensures users can only access their designated dashboard'),
    ('Responsive UI', 'Tailwind CSS ensures the application is accessible on desktop, tablet, and mobile'),
    ('RESTful APIs', 'Clean REST endpoints following HTTP standards (GET, POST, PUT, DELETE)'),
    ('Data Persistence', 'MySQL with JPA/Hibernate for automatic ORM and data persistence'),
    ('CORS Configuration', 'Spring Boot configured to accept requests from React frontend origin'),
]

for i, (feat, desc) in enumerate(features):
    row = feat_table.add_row()
    row.cells[0].text = feat
    row.cells[1].text = desc
    row.cells[0].paragraphs[0].runs[0].font.bold = True
    bg = 'ECEFF1' if i % 2 == 0 else 'FFFFFF'
    set_cell_bg(row.cells[0], bg)
    set_cell_bg(row.cells[1], bg)

# ============ SECTION 13: SECURITY FEATURES ============
doc.add_page_break()
add_heading(doc, '13. Security Features')
doc.add_paragraph(
    'Security is implemented at multiple layers of the application to protect user data and restrict '
    'unauthorized access:'
)

security_items = [
    ('Password Hashing (BCrypt)', 'All user passwords are hashed using BCrypt algorithm before storing in the database. Plain text passwords are never stored.'),
    ('Spring Security', 'Spring Security framework protects all backend API endpoints. Unauthorized requests receive 401/403 HTTP responses.'),
    ('Role-Based Authorization', 'APIs are secured by user role. FARMER role cannot access officer-specific endpoints and vice versa.'),
    ('JWT Token Authentication', 'After successful login, a JWT token is returned and stored in localStorage. This token is sent with every subsequent API request.'),
    ('CORS Configuration', 'Cross-Origin Resource Sharing is configured in Spring Boot to allow only the React frontend origin to communicate with the backend.'),
    ('Input Validation', 'Both frontend (React forms) and backend (Spring Boot validation annotations) validate user inputs to prevent invalid data.'),
    ('ProtectedRoute Component', "React Router's ProtectedRoute component checks authentication and role before rendering any protected page."),
    ('Stateless Architecture', 'The backend is stateless - no server-side sessions. Authentication is verified via token on each request.'),
]

sec_table = doc.add_table(rows=1, cols=2)
sec_table.style = 'Table Grid'
for i, h in enumerate(['Security Feature', 'Implementation Details']):
    sec_table.rows[0].cells[i].text = h
    sec_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
    set_cell_bg(sec_table.rows[0].cells[i], 'B71C1C')
    sec_table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

for i, (feat, desc) in enumerate(security_items):
    row = sec_table.add_row()
    row.cells[0].text = feat
    row.cells[1].text = desc
    row.cells[0].paragraphs[0].runs[0].font.bold = True
    bg = 'FFEBEE' if i % 2 == 0 else 'FFFFFF'
    set_cell_bg(row.cells[0], bg)
    set_cell_bg(row.cells[1], bg)

# ============ SECTION 14: CHALLENGES FACED ============
doc.add_page_break()
add_heading(doc, '14. Challenges Faced and Solutions')

challenges = [
    ('CORS Configuration', 'React (port 5173) and Spring Boot (port 8080) run on different ports, causing CORS errors when making API calls.', 'Added @CrossOrigin annotation on controllers and configured CorsRegistry in WebMvcConfigurer to allow frontend origin.'),
    ('Spring Security Setup', 'Configuring Spring Security to allow public routes (login, register) while protecting other endpoints was complex.', 'Created SecurityConfig class with custom filter chain, permitting /api/auth/** and securing /api/queries/** routes.'),
    ('Role-Based Routing (React)', 'Preventing unauthorized users from accessing dashboards by directly typing the URL.', 'Implemented ProtectedRoute component in React Router that checks localStorage for user role and redirects unauthorized access.'),
    ('JWT Token Integration', 'Attaching JWT token to every API request and handling token expiry.', 'Created Axios interceptors to automatically attach Authorization header to all requests from localStorage.'),
    ('JPA Entity Mapping', 'Mapping Java entity classes to MySQL tables with correct column names and constraints.', 'Used JPA annotations (@Entity, @Table, @Column, @Id) to map entities and Spring Data JPA for auto table creation.'),
    ('State Management in React', 'Managing authentication state and keeping UI in sync with backend data.', 'Used React hooks (useState, useEffect) for component state and re-fetching data after mutations.'),
    ('Password Security', 'Storing passwords securely in the database.', 'Used BCryptPasswordEncoder from Spring Security to hash passwords before saving and match during login.'),
]

chal_table = doc.add_table(rows=1, cols=3)
chal_table.style = 'Table Grid'
for i, h in enumerate(['Challenge', 'Problem', 'Solution']):
    chal_table.rows[0].cells[i].text = h
    chal_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
    set_cell_bg(chal_table.rows[0].cells[i], 'E65100')
    chal_table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

for i, (ch, prob, sol) in enumerate(challenges):
    row = chal_table.add_row()
    row.cells[0].text = ch
    row.cells[1].text = prob
    row.cells[2].text = sol
    row.cells[0].paragraphs[0].runs[0].font.bold = True
    bg = 'FFF3E0' if i % 2 == 0 else 'FFFFFF'
    for cell in row.cells:
        set_cell_bg(cell, bg)

# ============ SECTION 15: FUTURE ENHANCEMENTS ============
doc.add_page_break()
add_heading(doc, '15. Future Enhancements')
doc.add_paragraph(
    'While the current version of Farmers Buddy successfully implements core functionality, '
    'the following enhancements are planned for future versions:'
)

enhancements = [
    ('Push Notifications', 'Real-time notifications for farmers when officers reply to their queries, implemented using WebSockets or Firebase Cloud Messaging.'),
    ('Mobile Application', 'React Native mobile app for iOS and Android to make the platform accessible to farmers with smartphones.'),
    ('AI Crop Disease Detection', 'Integration of machine learning model to allow farmers to upload crop images and receive AI-powered disease diagnosis.'),
    ('Government Scheme Alerts', 'Automated notifications to farmers about relevant government agricultural schemes and subsidies.'),
    ('Multi-Language Support', 'Regional language support (Hindi, Marathi, Telugu, etc.) to make the platform accessible to farmers across India.'),
    ('SMS Alert Integration', 'SMS notifications using Twilio API for farmers who may not have consistent internet access.'),
    ('Admin Panel', 'Super-admin dashboard for system management, user management, and platform analytics.'),
    ('Analytics Dashboard', 'Visual charts and statistics for officers showing query trends, response times, and common query topics.'),
    ('Document Upload', 'Allow farmers to attach photos of crops or documents with their queries for better context.'),
    ('Rating System', 'Allow farmers to rate the helpfulness of officer replies for quality monitoring.'),
]

enh_table = doc.add_table(rows=1, cols=3)
enh_table.style = 'Table Grid'
for i, h in enumerate(['#', 'Enhancement', 'Description']):
    enh_table.rows[0].cells[i].text = h
    enh_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
    set_cell_bg(enh_table.rows[0].cells[i], '1A237E')
    enh_table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

for i, (enh, desc) in enumerate(enhancements):
    row = enh_table.add_row()
    row.cells[0].text = str(i + 1)
    row.cells[1].text = enh
    row.cells[2].text = desc
    row.cells[1].paragraphs[0].runs[0].font.bold = True
    bg = 'E8EAF6' if i % 2 == 0 else 'FFFFFF'
    for cell in row.cells:
        set_cell_bg(cell, bg)

# ============ SECTION 16: CONCLUSION ============
doc.add_page_break()
add_heading(doc, '16. Conclusion')

doc.add_paragraph(
    'Farmers Buddy is a fully functional full-stack web application that successfully demonstrates '
    'the integration of modern frontend and backend technologies to solve a real-world agricultural '
    'communication problem.'
)
doc.add_paragraph(
    'The project implements a complete authentication and authorization system using Spring Security '
    'and BCrypt, a role-based dashboard system differentiating Farmer and Officer experiences, and a '
    'clean RESTful API architecture connecting a React.js frontend with a Spring Boot backend and '
    'MySQL database.'
)
doc.add_paragraph(
    'The application is built following software engineering best practices - clean separation of '
    'concerns through layered architecture, stateless authentication with JWT tokens, ORM-based '
    'database operations with JPA/Hibernate, and responsive UI design with Tailwind CSS.'
)
doc.add_paragraph(
    'This project provided comprehensive hands-on experience with the complete modern web development '
    'stack: React.js for dynamic UI, Spring Boot for robust backend APIs, Spring Security for '
    'authentication, MySQL for data persistence, and GitHub for version control. The skills and '
    'patterns learned through this project are directly applicable to industry-level software development.'
)

# Final viva summary table
doc.add_paragraph()
add_heading(doc, 'Quick Viva Reference Summary', level=2)

summary_table = doc.add_table(rows=1, cols=2)
summary_table.style = 'Table Grid'
for i, h in enumerate(['Question', 'Answer']):
    summary_table.rows[0].cells[i].text = h
    summary_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
    set_cell_bg(summary_table.rows[0].cells[i], '2E7D32')
    summary_table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

viva_qa = [
    ('What is Farmers Buddy?', 'A full-stack agriculture support platform for farmer-officer query communication with role-based dashboards.'),
    ('What is the frontend technology?', 'React.js with Tailwind CSS for styling, Axios for HTTP calls, and React Router for navigation.'),
    ('What is the backend technology?', 'Spring Boot with Spring Security and Spring Data JPA/Hibernate.'),
    ('What database is used?', 'MySQL - with two main tables: USERS and QUERIES.'),
    ('How are passwords stored?', "BCrypt hashed - Spring Security's BCryptPasswordEncoder hashes passwords before saving."),
    ('How is authentication handled?', 'JWT token-based authentication. Token stored in localStorage after login.'),
    ('What is role-based access?', 'FARMER role -> Farmer Dashboard. OFFICER role -> Officer Dashboard. Protected by ProtectedRoute in React and Spring Security.'),
    ('What does the architecture look like?', 'React <-> REST API <-> Spring Boot Controllers <-> Services <-> JPA Repositories <-> MySQL.'),
    ('What is JPA/Hibernate?', 'Java Persistence API - maps Java classes (entities) to database tables automatically without writing SQL.'),
    ('What was the main challenge?', 'CORS configuration between React (port 5173) and Spring Boot (port 8080) running on different ports.'),
]

for i, (q, a) in enumerate(viva_qa):
    row = summary_table.add_row()
    row.cells[0].text = q
    row.cells[1].text = a
    row.cells[0].paragraphs[0].runs[0].font.bold = True
    bg = 'E8F5E9' if i % 2 == 0 else 'FFFFFF'
    for cell in row.cells:
        set_cell_bg(cell, bg)

# Save the document
doc.save('/Users/I578365/farmers-buddy/Farmers_Buddy_Project_Report.docx')
print('Document saved successfully!')
